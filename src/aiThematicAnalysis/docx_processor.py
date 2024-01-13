#! /usr/bin/env python3
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
import rich_click as click
from rich.console import Console
from pathlib import Path
import pandas as pd


def docInit():
    docTitle = 'Thematic Analysis'
    # Create the document
    document = Document()
    style = document.styles['Normal']
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Info Document
    document.core_properties.author = "Enzo Mazzeo"
    document.core_properties.title = docTitle
    document.core_properties.version = 1
    document.core_properties.revision = 1
    document.core_properties.keywords = "Thematic analysis, Reports"

    document.add_heading(docTitle, 0)
    paragraph1 = document.add_paragraph(
        "This document contain the table from the AI response.")

    return document


@click.command()
@click.option('-f', '--folder', metavar="FOLDER", help="Folder to anlyze")
def action(folder):
    console = Console()
    folder = Path(folder)
    document = docInit()

    themes = folder.joinpath('results-ph3.csv')
    if themes.exists():
        df2 = pd.read_csv(themes).fillna('')
    else:
        console.print(f"The file {themes} not exists")
        exit(1)
    table = document.add_table(
        rows=len(df2)+1, cols=4, style="Colorful List Accent 6")
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Nr.'
    hdr_cells[1].text = 'Nome tema'
    hdr_cells[2].text = "Descrizione"
    hdr_cells[3].text = "Topics"
    for index, row in df2.iterrows():
        text_cells = table.rows[index+1].cells
        text_cells[0].text = f"{index+1}"
        text_cells[1].text = f"{row['Group']}"
        text_cells[2].text = f"{row['description']}"
        text_cells[3].text = f"{row['topics']}"
    document.save(folder.joinpath('results-ph3.docx'))


if __name__ == "__main__":
    action()
